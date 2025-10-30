import PyPDF2
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import nltk
from io import BytesIO
import json
import uuid
import tempfile
import os

class DocumentProcessor:
    def __init__(self):
        self.figures = {}
        self.tables = {}
        self.processing_errors = []
        
    def extract_docx_elements(self, file_path):
        """Extract tables and figures from DOCX with error tracking"""
        try:
            doc = Document(file_path)
            clean_text = ""
            
            for i, paragraph in enumerate(doc.paragraphs):
                # Check for inline shapes (figures)
                if paragraph._element.xpath('.//pic:pic'):
                    figure_id = f"FIGURE_{len(self.figures)+1:03d}"
                    try:
                        # Extract figure info
                        self.figures[figure_id] = {
                            'type': 'image',
                            'caption': paragraph.text.strip(),
                            'position': i,
                            'status': 'extracted'
                        }
                        # Add placeholder
                        clean_text += f"\n[{figure_id}]\n"
                    except Exception as e:
                        self.processing_errors.append(f"Figure {len(self.figures)+1}: {str(e)}")
                        clean_text += paragraph.text + "\n"
                else:
                    clean_text += paragraph.text + "\n"
            
            # Extract tables
            for table_idx, table in enumerate(doc.tables):
                table_id = f"TABLE_{table_idx+1:03d}"
                try:
                    # Extract table data
                    table_data = []
                    for row_idx, row in enumerate(table.rows):
                        row_data = []
                        for cell in row.cells:
                            row_data.append(cell.text.strip())
                        table_data.append(row_data)
                    
                    self.tables[table_id] = {
                        'data': table_data,
                        'rows': len(table_data),
                        'cols': len(table_data[0]) if table_data else 0,
                        'status': 'extracted'
                    }
                    
                    # Add placeholder in text
                    clean_text = clean_text.replace(
                        self._get_table_text(table), 
                        f"\n[{table_id}]\n"
                    )
                    
                except Exception as e:
                    self.processing_errors.append(f"Table {table_idx+1}: {str(e)}")
                    self.tables[table_id] = {'status': 'failed', 'error': str(e)}
            
            return clean_text
            
        except Exception as e:
            self.processing_errors.append(f"Document processing: {str(e)}")
            return None
    
    def _get_table_text(self, table):
        """Extract text representation of table for replacement"""
        text = ""
        for row in table.rows:
            for cell in row.cells:
                text += cell.text + " "
            text += "\n"
        return text.strip()
    
    def translate_table_cells(self, table_id, translate_function):
        """Translate table cells with error tracking"""
        if table_id not in self.tables:
            self.processing_errors.append(f"{table_id}: Table not found")
            return False
        
        table = self.tables[table_id]
        if table['status'] != 'extracted':
            return False
        
        try:
            translated_data = []
            failed_cells = []
            
            for row_idx, row in enumerate(table['data']):
                translated_row = []
                for col_idx, cell_text in enumerate(row):
                    if cell_text.strip():
                        try:
                            translated_cell = translate_function(cell_text)
                            if translated_cell.startswith("Error:"):
                                failed_cells.append(f"Row {row_idx+1}, Col {col_idx+1}")
                                translated_cell = cell_text  # Keep original on error
                            translated_row.append(translated_cell)
                        except Exception as e:
                            failed_cells.append(f"Row {row_idx+1}, Col {col_idx+1}: {str(e)}")
                            translated_row.append(cell_text)
                    else:
                        translated_row.append(cell_text)
                translated_data.append(translated_row)
            
            if failed_cells:
                self.processing_errors.append(f"{table_id} - Failed cells: {', '.join(failed_cells)}")
                table['status'] = 'partial'
            else:
                table['status'] = 'translated'
            
            table['translated_data'] = translated_data
            return len(failed_cells) == 0
            
        except Exception as e:
            self.processing_errors.append(f"{table_id}: Translation failed - {str(e)}")
            table['status'] = 'failed'
            return False
    
    def reconstruct_docx(self, translated_text, output_path):
        """Reconstruct DOCX with translated content and elements"""
        try:
            import re
            doc = Document()
            
            # Split by double newlines to preserve paragraph structure
            paragraphs = translated_text.split('\n\n')
            
            for para in paragraphs:
                para = para.strip()
                if not para:
                    continue
                
                # Check if this paragraph contains only a table placeholder
                table_match = re.fullmatch(r'\[TABLE_\d{3}\]', para)
                if table_match:
                    table_id = para[1:-1]  # Remove brackets
                    if not self._insert_table(doc, table_id):
                        self.processing_errors.append(f"{table_id}: Failed to insert")
                        doc.add_paragraph(f"[{table_id} - Failed to insert]")
                    continue
                
                # Check if this paragraph contains only a figure placeholder
                figure_match = re.fullmatch(r'\[FIGURE_\d{3}\]', para)
                if figure_match:
                    figure_id = para[1:-1]
                    if not self._insert_figure_placeholder(doc, figure_id):
                        self.processing_errors.append(f"{figure_id}: Failed to insert")
                        doc.add_paragraph(f"[{figure_id} - Failed to insert]")
                    continue
                
                # Check if paragraph contains placeholders mixed with text
                if '[TABLE_' in para or '[FIGURE_' in para:
                    # Split by placeholders and process each part
                    parts = re.split(r'(\[(?:TABLE|FIGURE)_\d{3}\])', para)
                    
                    for part in parts:
                        part = part.strip()
                        if not part:
                            continue
                        
                        # Check if it's a table placeholder
                        if re.fullmatch(r'\[TABLE_\d{3}\]', part):
                            table_id = part[1:-1]
                            if not self._insert_table(doc, table_id):
                                doc.add_paragraph(f"[{table_id} - Failed]")
                        
                        # Check if it's a figure placeholder
                        elif re.fullmatch(r'\[FIGURE_\d{3}\]', part):
                            figure_id = part[1:-1]
                            if not self._insert_figure_placeholder(doc, figure_id):
                                doc.add_paragraph(f"[{figure_id} - Failed]")
                        
                        # Regular text
                        else:
                            doc.add_paragraph(part)
                else:
                    # Regular paragraph with no placeholders
                    doc.add_paragraph(para)
            
            doc.save(output_path)
            return True
            
        except Exception as e:
            self.processing_errors.append(f"Document reconstruction: {str(e)}")
            return False
    
    def _insert_table(self, doc, table_id):
        """Insert translated table into document"""
        if table_id not in self.tables:
            return False
        
        table_info = self.tables[table_id]
        
        # Use translated data if available, otherwise use original data
        if 'translated_data' in table_info:
            data = table_info['translated_data']
        elif 'data' in table_info:
            data = table_info['data']
        else:
            return False
        
        try:
            if not data or len(data) == 0:
                return False
                
            table = doc.add_table(rows=len(data), cols=len(data[0]))
            table.style = 'Table Grid'
            
            for row_idx, row_data in enumerate(data):
                for col_idx, cell_text in enumerate(row_data):
                    table.cell(row_idx, col_idx).text = str(cell_text)
            
            return True
        except Exception as e:
            self.processing_errors.append(f"{table_id} insertion error: {str(e)}")
            return False
    
    def _insert_figure_placeholder(self, doc, figure_id):
        """Insert figure placeholder into document"""
        if figure_id not in self.figures:
            return False
        
        try:
            figure_info = self.figures[figure_id]
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.add_run(f"[{figure_id}: {figure_info.get('caption', 'Image')}]").bold = True
            return True
        except Exception:
            return False
    
    def get_processing_summary(self):
        """Get summary of processing results"""
        total_tables = len(self.tables)
        total_figures = len(self.figures)
        
        successful_tables = sum(1 for t in self.tables.values() if t['status'] == 'translated')
        successful_figures = sum(1 for f in self.figures.values() if f['status'] == 'extracted')
        
        return {
            'tables': {'total': total_tables, 'successful': successful_tables},
            'figures': {'total': total_figures, 'successful': successful_figures},
            'errors': self.processing_errors
        }